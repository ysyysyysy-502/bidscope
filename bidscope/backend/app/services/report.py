import re
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from app.core.config import settings
from app.domain.enums import SourceTier
from app.domain.schemas import Notice, QueryIntent, SourcePlan

INVALID = r'[\\/:*?"<>|\n\r\t]+'

def safe_filename(query: str, now: datetime | None = None) -> str:
    now = now or datetime.now()
    clean = re.sub(INVALID, "_", query).strip(" _")
    clean = clean[:70] or "标讯查询"
    return f"{clean}_{now.strftime('%Y%m%d%H%M')}.docx"

def _add_kv(table, key, value):
    row = table.add_row().cells
    row[0].text = key
    row[1].text = value or "—"

def generate_word(intent: QueryIntent, plans: list[SourcePlan], notices: list[Notice], duplicate_count: int, skipped_incremental: int) -> Path:
    doc = Document()
    styles = doc.styles
    styles['Normal'].font.name = 'Microsoft YaHei'
    styles['Normal'].font.size = Pt(10.5)
    title = doc.add_heading('标讯罗盘 BidScope 标讯汇总报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('本报告由 Demo 系统生成，样例数据不代表真实外部站点采集结果。').italic = True

    doc.add_heading('一、查询信息', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '字段'
    table.rows[0].cells[1].text = '内容'
    _add_kv(table, '用户问题', intent.raw_query)
    _add_kv(table, '主题', intent.topic)
    _add_kv(table, '同义词', '、'.join(intent.synonyms))
    _add_kv(table, '地区', '、'.join(intent.regions))
    _add_kv(table, '时间范围', f"{intent.start_at:%Y-%m-%d} 至 {intent.end_at:%Y-%m-%d}")
    _add_kv(table, '频率', intent.schedule_text or '一次性查询')
    _add_kv(table, '意图置信度', str(intent.confidence))

    doc.add_heading('二、来源计划与权限边界', level=1)
    st = doc.add_table(rows=1, cols=5)
    st.style = 'Table Grid'
    headers = ['来源', '证据等级', '访问状态', '模式', '说明']
    for i,h in enumerate(headers): st.rows[0].cells[i].text = h
    for plan in plans:
        cells = st.add_row().cells
        cells[0].text = plan.source_name
        cells[1].text = plan.tier.value
        cells[2].text = plan.access_status.value
        cells[3].text = plan.mode
        cells[4].text = plan.note

    formal = [n for n in notices if n.source_tier == SourceTier.T0_CANONICAL]
    signals = [n for n in notices if n.source_tier != SourceTier.T0_CANONICAL]

    doc.add_heading('三、结果统计', level=1)
    doc.add_paragraph(f'本次新增可交付结果：{len(notices)} 条；其中正式公告 {len(formal)} 条，待核验线索 {len(signals)} 条。')
    doc.add_paragraph(f'确定性去重过滤：{duplicate_count} 条；增量账本跳过历史已推送：{skipped_incremental} 条。')
    restricted = sum(1 for n in notices if n.restricted_fields)
    doc.add_paragraph(f'存在受限字段的商业源记录：{restricted} 条。系统未尝试恢复遮罩、会员或无权访问字段。')

    def add_notice_section(title, items):
        doc.add_heading(title, level=1)
        if not items:
            doc.add_paragraph('无。')
            return
        for idx, n in enumerate(items, 1):
            doc.add_heading(f'{idx}. {n.title}', level=2)
            t = doc.add_table(rows=1, cols=2)
            t.style = 'Table Grid'
            t.rows[0].cells[0].text = '字段'
            t.rows[0].cells[1].text = '内容'
            _add_kv(t, '标题', n.title)
            _add_kv(t, '发布时间', n.published_at.strftime('%Y-%m-%d %H:%M'))
            _add_kv(t, '来源链接', n.source_url)
            _add_kv(t, '核心内容', n.core_content)
            _add_kv(t, '附件链接', '\n'.join([f'{a.name}: {a.url}' for a in n.attachments]) if n.attachments else '无附件')
            _add_kv(t, '采购人', n.buyer)
            _add_kv(t, '项目编号', n.project_no)
            _add_kv(t, '公告阶段', n.notice_stage.value)
            _add_kv(t, '预算/最高限价', n.budget)
            _add_kv(t, '证据等级', n.source_tier.value)
            _add_kv(t, '命中理由', n.relevance_reason)
            _add_kv(t, '证据片段', '\n'.join([e.quote for e in n.evidence]) if n.evidence else '无')
            if n.restricted_fields:
                _add_kv(t, '受限字段', '、'.join(n.restricted_fields))
    add_notice_section('四、已回源核验的正式公告', formal)
    add_notice_section('五、待核验机会信号', signals)

    doc.add_heading('六、失败、受限与真实性说明', level=1)
    doc.add_paragraph('1. 本 Demo 默认使用脱敏样例数据与模拟连接器，不代表已经真实接入中国政府采购网、商业平台或 API 服务。')
    doc.add_paragraph('2. 遇到验证码、付费墙、权限不足或明确禁止自动化时，系统应暂停并回链人工处理。')
    doc.add_paragraph('3. AI 只生成相关性理由与摘要，不改写标题、发布时间、链接、预算、截止时间等事实字段。')
    doc.add_paragraph('4. 正式 POC 需使用人工金标准验证准确率、召回率、误合并率、附件召回和重复推送。')

    filename = safe_filename(intent.raw_query)
    path = settings.storage_dir / filename
    doc.save(path)
    return path
